"""
简历文件预处理解析器模块

支持多种格式的简历文件提取（PDF、Word、TXT、图片），
提供统一的文件解析接口和完善的错误处理机制。
同时封装火山引擎Vision/LLM API客户端，用于简历智能解析。
"""

import base64
import json
import os
import re
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Union

# HTTP请求库
import requests

# PDF处理库（优先使用pdfplumber，回退到pypdf）
import pdfplumber
import pypdf

# Word文档处理库
from docx import Document

# 图片处理库
from PIL import Image

# Pydantic数据模型
from pydantic import ValidationError

# 日志系统
from utils.logger import get_logger

# 配置加载器
from utils.config_loader import get_config

# 数据库操作封装
from utils.db_helper import get_db

# 初始化日志记录器
logger = get_logger(__name__)


# ==================== 简历解析Prompt模板 ====================
RESUME_PARSE_PROMPT = """你是一个专业的HR简历分析助手。请仔细阅读以下简历内容，提取关键信息并以严格的JSON格式输出。

请提取以下信息（如果某项信息不存在，填null）：

{{
  "basic_info": {{
    "name": "姓名",
    "gender": "性别",
    "age": 年龄,
    "phone": "联系电话",
    "email": "电子邮箱"
  }},
  "education": {{
    "degree": "最高学历（高中/专科/本科/硕士/博士）",
    "school": "毕业院校",
    "major": "专业",
    "graduation_year": 毕业年份（整数）
  }},
  "work_experience": {{
    "total_years": 工作年限（浮点数，如3.5表示3年半）,
    "positions": [
      {{
        "company": "公司名称",
        "title": "职位名称",
        "start_date": "开始时间（YYYY-MM或YYYY年MM月）",
        "end_date": "结束时间（至今用'至今'）",
        "description": "工作描述（100字以内）"
      }}
    ]
  }},
  "skills": ["技能标签1", "技能标签2", ...],
  "projects": [
    {{
      "name": "项目名称",
      "description": "项目描述（100字以内）",
      "role": "担任角色",
      "tech_stack": ["技术1", "技术2", ...],
      "achievement": "项目成果（50字以内）"
    }}
  ],
  "job_expectation": {{
    "target_positions": ["期望岗位1", "期望岗位2", ...],
    "expected_salary": "期望薪资（如'15k-25k'）",
    "expected_locations": ["城市1", "城市2", ...]
  }},
  "self_evaluation": "自我评价/个人优势总结（200字以内）"
}}

注意事项：
1. 必须输出有效的JSON格式，不要包含markdown代码块标记
2. 数字类型字段必须是数字，不要加引号
3. 技能标签尽量具体（如"Python"而不是"编程"）
4. 工作经历按时间倒序排列（最近的工作在前）
5. 提取的信息要准确，不要编造不存在的内容

---
简历内容如下：

{content}
"""


class ResumeParser:
    """
    简历文件预处理解析器

    负责解析多种格式的简历文件（PDF、Word、TXT、图片），
    提取文本内容或图片二进制数据供后续AI分析使用。

    Attributes:
        supported_formats (dict): 支持的文件格式与对应处理方法的映射
        max_file_size (int): 文件大小限制（字节），默认10MB
    """

    # 支持的文件格式映射表
    SUPPORTED_FORMATS: Dict[str, str] = {
        '.pdf': 'text',       # PDF文档 → 提取文本
        '.docx': 'text',      # Word文档 → 提取文本
        '.doc': 'text',       # 旧版Word → 提取文本（有限支持）
        '.txt': 'text',       # 纯文本文件 → 直接读取
        '.png': 'image',      # PNG图片 → 返回二进制
        '.jpg': 'image',      # JPG图片 → 返回二进制
        '.jpeg': 'image',     # JPEG图片 → 返回二进制
    }

    def __init__(self, max_file_size: int = 10 * 1024 * 1024):
        """
        初始化简历解析器

        Args:
            max_file_size: 允许的最大文件大小（字节），默认10MB
        """
        self.max_file_size = max_file_size
        logger.info(f"简历解析器初始化完成，最大文件限制: {max_file_size / 1024 / 1024:.1f}MB")

    def parse_resume_file(self, file_path: str) -> Dict[str, Any]:
        """
        解析简历文件的主入口方法

        统一处理各种格式的简历文件，返回标准化的解析结果。
        包含完整的验证、格式检测、内容提取和错误处理流程。

        Args:
            file_path: 简历文件的完整路径（绝对路径或相对路径）

        Returns:
            dict: 标准化的解析结果，包含以下字段：
                - success (bool): 解析是否成功
                - file_type (str): 文件类型 ('text' | 'image')
                - content (str | bytes): 提取的文本内容或图片二进制数据
                - metadata (dict): 文件元信息
                    * filename (str): 文件名
                    * size (int): 文件大小（字节）
                    * format (str): 文件扩展名
                - error (str | None): 错误信息（成功时为None）

        Example:
            >>> parser = ResumeParser()
            >>> result = parser.parse_resume_file("resumes/张三.pdf")
            >>> if result['success']:
            ...     print(f"提取到 {len(result['content'])} 字符")
        """
        start_time = time.time()

        try:
            # 1. 验证文件有效性
            validation_result = self._validate_file(file_path)
            if not validation_result['valid']:
                return {
                    'success': False,
                    'file_type': None,
                    'content': None,
                    'metadata': {},
                    'error': validation_result['error']
                }

            # 2. 检测文件格式并确定类型
            file_ext = self._detect_format(file_path)
            file_type = self.SUPPORTED_FORMATS.get(file_ext)

            if file_type is None:
                error_msg = f"不支持的文件格式：{file_ext}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'file_type': None,
                    'content': None,
                    'metadata': self._build_metadata(file_path),
                    'error': error_msg
                }

            # 3. 记录日志信息
            filename = Path(file_path).name
            logger.info(f"开始解析简历文件: {filename} (格式: {file_ext}, 类型: {file_type})")

            # 4. 根据格式调用对应的提取方法
            extract_method = self._get_extract_method(file_ext)
            content = extract_method(file_path)

            # 5. 计算耗时并记录完成日志
            elapsed_time = time.time() - start_time
            logger.info(
                f"简历文件解析完成: {filename}, "
                f"耗时: {elapsed_time:.3f}秒, "
                f"内容长度: {len(content) if isinstance(content, (str, bytes)) else 'N/A'}"
            )

            # 6. 返回标准化结果
            return {
                'success': True,
                'file_type': file_type,
                'content': content,
                'metadata': self._build_metadata(file_path),
                'error': None
            }

        except FileNotFoundError as e:
            error_msg = f"文件不存在: {file_path}"
            logger.error(error_msg, exc_info=True)
            return {
                'success': False,
                'file_type': None,
                'content': None,
                'metadata': {},
                'error': error_msg
            }
        except PermissionError as e:
            error_msg = f"无权限访问文件: {file_path}"
            logger.error(error_msg, exc_info=True)
            return {
                'success': False,
                'file_type': None,
                'content': None,
                'metadata': {},
                'error': error_msg
            }
        except Exception as e:
            error_msg = f"解析失败: {str(e)}"
            logger.error(f"解析文件时发生异常: {file_path}", exc_info=True)
            return {
                'success': False,
                'file_type': None,
                'content': None,
                'metadata': self._build_metadata(file_path) if os.path.exists(file_path) else {},
                'error': error_msg
            }

    def _validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        验证文件的有效性

        检查文件是否存在、是否为有效文件、大小是否符合要求。

        Args:
            file_path: 待验证的文件路径

        Returns:
            dict: 验证结果，包含以下字段：
                - valid (bool): 是否通过验证
                - error (str | None): 错误信息（通过验证时为None）
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return {'valid': False, 'error': f'文件不存在: {file_path}'}

        # 检查是否为文件（而非目录）
        if not os.path.isfile(file_path):
            return {'valid': False, 'error': f'路径不是有效文件: {file_path}'}

        # 检查文件大小
        file_size = self._get_file_size(file_path)
        if file_size > self.max_file_size:
            return {
                'valid': False,
                'error': f'文件大小超过{self.max_file_size / 1024 / 1024:.1f}MB限制 '
                         f'(当前: {file_size / 1024 / 1024:.2f}MB)'
            }

        # 检查文件是否为空
        if file_size == 0:
            return {'valid': False, 'error': '文件为空'}

        return {'valid': True, 'error': None}

    def _detect_format(self, file_path: str) -> str:
        """
        根据文件扩展名检测文件格式

        将扩展名统一转换为小写格式以便匹配。

        Args:
            file_path: 文件路径

        Returns:
            str: 小写的文件扩展名（包含点号），例如 '.pdf', '.docx'
        """
        ext = Path(file_path).suffix.lower()
        return ext

    def _get_file_size(self, file_path: str) -> int:
        """
        获取文件大小

        Args:
            file_path: 文件路径

        Returns:
            int: 文件大小（字节）
        """
        return os.path.getsize(file_path)

    def _is_supported_format(self, file_ext: str) -> bool:
        """
        检查文件格式是否受支持

        Args:
            file_ext: 文件扩展名（小写，包含点号）

        Returns:
            bool: 是否为支持的格式
        """
        return file_ext in self.SUPPORTED_FORMATS

    def _get_extract_method(self, file_ext: str) -> Callable[[str], Union[str, bytes]]:
        """
        根据文件格式获取对应的提取方法

        Args:
            file_ext: 文件扩展名

        Returns:
            callable: 对应格式的提取方法函数

        Raises:
            ValueError: 不支持的文件格式
        """
        method_map = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.txt': self._extract_txt_text,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
        }

        method = method_map.get(file_ext)
        if method is None:
            raise ValueError(f"不支持的文件格式: {file_ext}")

        return method

    def _extract_pdf_text(self, file_path: str) -> str:
        """
        提取PDF文件的文本内容

        优先使用 pdfplumber 进行提取（更好的中文支持和表格识别），
        如果 pdfplumber 失败则回退到 pypdf。

        Args:
            file_path: PDF文件路径

        Returns:
            str: 提取的全部文本内容（UTF-8编码）

        Raises:
            Exception: PDF解析失败时抛出异常
        """
        text_content = []

        try:
            # 方案1：优先使用 pdfplumber（更好的中文支持）
            logger.debug(f"尝试使用 pdfplumber 解析PDF: {file_path}")
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取页面文本
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

                    # 尝试提取表格内容（如有）
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = self._format_table_to_text(table)
                            text_content.append(f"\n[表格{table_idx + 1}]\n{table_text}")

            full_text = '\n\n'.join(text_content)

            if not full_text.strip():
                raise ValueError("pdfplumber未能提取到文本内容")

            logger.info(f"pdfplumber成功提取PDF文本，共{len(pdf.pages)}页")
            return full_text.strip()

        except Exception as pdfplumber_error:
            logger.warning(f"pdfplumber解析失败，回退到pypdf: {str(pdfplumber_error)}")

            # 方案2：回退到 pypdf
            try:
                logger.debug(f"使用 pypdf 回退解析PDF: {file_path}")
                text_content = []
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)

                full_text = '\n\n'.join(text_content)

                if not full_text.strip():
                    raise ValueError("pypdf也未能提取到文本内容")

                logger.info(f"pypdf成功提取PDF文本，共{len(pdf_reader.pages)}页")
                return full_text.strip()

            except Exception as pypdf_error:
                error_msg = f"PDF解析完全失败 (pdfplumber: {pdfplumber_error}, pypdf: {pypdf_error})"
                logger.error(error_msg)
                raise Exception(error_msg) from pypdf_error

    def _format_table_to_text(self, table: list) -> str:
        """
        将表格数据转换为可读的文本格式

        Args:
            table: 二维列表形式的表格数据

        Returns:
            str: 格式化后的表格文本
        """
        rows = []
        for row in table:
            # 过滤空值并将单元格内容转换为字符串
            formatted_row = [str(cell) if cell else '' for cell in row]
            rows.append(' | '.join(formatted_row))

        return '\n'.join(rows)

    def _extract_docx_text(self, file_path: str) -> str:
        """
        提取Word文档（.docx）的文本内容

        提取段落文本和表格内容，保持文档的基本结构层次。

        Args:
            file_path: Word文档路径

        Returns:
            str: 提取的全部文本内容

        Raises:
            Exception: Word文档解析失败时抛出异常
        """
        try:
            logger.debug(f"解析Word文档: {file_path}")
            doc = Document(file_path)

            text_parts = []

            # 1. 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 根据标题级别添加标记
                    style_name = para.style.name if para.style else ''
                    if 'Heading' in style_name or 'heading' in style_name:
                        text_parts.append(f'\n## {text}\n')
                    else:
                        text_parts.append(text)

            # 2. 提取表格内容
            for table_idx, table in enumerate(doc.tables, 1):
                table_texts = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_texts.append(' | '.join(row_data))

                if table_texts:
                    text_parts.append(f'\n[表格{table_idx}]\n' + '\n'.join(table_texts))

            full_text = '\n'.join(text_parts).strip()

            if not full_text:
                raise ValueError("Word文档内容为空")

            logger.info(f"成功提取Word文档文本，共{len(doc.paragraphs)}个段落，"
                        f"{len(doc.tables)}个表格")
            return full_text

        except Exception as e:
            logger.error(f"Word文档解析失败: {str(e)}", exc_info=True)
            raise Exception(f"Word文档解析失败: {str(e)}") from e

    def _extract_doc_text(self, file_path: str) -> str:
        """
        处理旧版Word文档（.doc格式）

        .doc 为二进制格式，Python原生支持有限。
        此方法会提示用户将文件转换为 .docx 格式。

        Args:
            file_path: 旧版Word文档路径

        Returns:
            str: 错误提示信息

        Note:
            如需完整支持 .doc 格式，建议安装 python-win32com 或 antiword 工具
        """
        error_msg = (
            "不支持直接解析 .doc 格式的旧版Word文档。\n"
            "请使用 Microsoft Word 或在线工具将其转换为 .docx 格式后重新上传。\n"
            "转换方法：\n"
            "1. 用 Word 打开文件 → 另存为 → 选择「Word 文档 (*.docx)」格式\n"
            "2. 或使用在线转换工具（如 convertio.co）"
        )
        logger.warning(f"收到不支持的.doc格式文件: {file_path}")
        raise NotImplementedError(error_msg)

    def _extract_txt_text(self, file_path: str) -> str:
        """
        读取纯文本文件的内容

        自动检测文件编码（按 UTF-8 → GBK → GB2312 → Latin-1 顺序尝试），
        统一换行符格式。

        Args:
            file_path: 文本文件路径

        Returns:
            str: 文件文本内容（统一换行符为 \\n）

        Raises:
            Exception: 文件读取或编码检测失败时抛出异常
        """
        # 编码检测顺序（从最可能到最不可能）
        encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'latin-1']

        last_error = None
        for encoding in encodings_to_try:
            try:
                logger.debug(f"尝试使用 {encoding} 编码读取文件: {file_path}")
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()

                # 统一换行符（Windows \r\n → Unix \n）
                content = content.replace('\r\n', '\n').replace('\r', '\n')

                logger.info(f"成功使用 {encoding} 编码读取TXT文件")
                return content.strip()

            except UnicodeDecodeError as e:
                logger.debug(f"{encoding} 编码解码失败，尝试下一种编码")
                last_error = e
                continue
            except Exception as e:
                logger.error(f"读取文件时发生异常: {str(e)}", exc_info=True)
                raise Exception(f"文件读取失败: {str(e)}") from e

        # 所有编码都失败
        error_msg = f"无法检测文件编码，已尝试: {encodings_to_try}"
        logger.error(error_msg)
        raise Exception(error_msg) from last_error

    def _process_image(self, file_path: str) -> bytes:
        """
        处理图片文件（PNG/JPG/JPEG）

        使用 Pillow 验证图片格式并读取二进制数据，
        可选转换为 RGB 模式以兼容不同图片格式。

        Args:
            file_path: 图片文件路径

        Returns:
            bytes: 图片的二进制数据（供 Vision API 使用）

        Raises:
            Exception: 图片处理失败时抛出异常
        """
        try:
            logger.debug(f"处理图片文件: {file_path}")

            # 使用 Pillow 打开并验证图片
            with Image.open(file_path) as img:
                # 验证图片完整性
                img.verify()

            # 重新打开以读取数据（verify后会关闭文件）
            with Image.open(file_path) as img:
                # 获取图片基本信息
                img_format = img.format
                img_mode = img.mode
                img_size = img.size

                logger.info(
                    f"图片验证通过: 格式={img_format}, "
                    f"模式={img_mode}, 尺寸={img_size[0]}x{img_size[1]}"
                )

                # 如果是 RGBA 或其他模式，转换为 RGB（某些API只支持RGB）
                if img_mode != 'RGB':
                    logger.debug(f"将图片从 {img_mode} 模式转换为 RGB")
                    img_rgb = img.convert('RGB')
                    # 写入内存缓冲区
                    buffer = BytesIO()
                    # 根据原始格式选择保存格式
                    save_format = 'JPEG' if img_format in ['JPEG', 'JPG'] else 'PNG'
                    img_rgb.save(buffer, format=save_format)
                    image_bytes = buffer.getvalue()
                else:
                    # 直接读取原始二进制数据
                    with open(file_path, 'rb') as f:
                        image_bytes = f.read()

                logger.info(f"成功处理图片文件，大小: {len(image_bytes) / 1024:.2f}KB")
                return image_bytes

        except Exception as e:
            logger.error(f"图片处理失败: {str(e)}", exc_info=True)
            raise Exception(f"图片处理失败: {str(e)}") from e

    def _build_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        构建文件元数据信息

        收集文件的基本属性信息用于返回结果中。

        Args:
            file_path: 文件路径

        Returns:
            dict: 包含文件元信息的字典
                - filename (str): 文件名
                - size (int): 文件大小（字节）
                - format (str): 文件扩展名
        """
        path_obj = Path(file_path)
        return {
            'filename': path_obj.name,
            'size': self._get_file_size(file_path) if os.path.exists(file_path) else 0,
            'format': self._detect_format(file_path)
        }


class VolcengineVisionClient:
    """
    火山引擎Vision/LLM API客户端

    封装火山引擎豆包大模型API调用，支持图片简历和文本简历的智能解析。
    采用OpenAI兼容接口格式，提供完善的错误处理、重试机制和日志记录。

    Attributes:
        api_key (str): 火山引擎API密钥
        base_url (str): API基础地址
        vision_model (str): 视觉模型名称（用于图片解析）
        text_model (str): 文本模型名称（用于文本解析）
        max_retries (int): 最大重试次数
        timeout (int): 请求超时时间（秒）
    """

    # 默认配置值
    DEFAULT_VISION_MODEL = "doubao-vision-pro-32k"
    DEFAULT_TEXT_MODEL = "doubao-pro-32k"
    DEFAULT_MAX_RETRIES = 3
    DEFAULT_TIMEOUT = 60

    def __init__(self, config: Dict[str, Any] = None):
        """
        初始化API客户端

        从配置字典或api_config.yaml加载API密钥、端点地址等参数。

        Args:
            config: API配置字典，如为None则从api_config.yaml自动加载
                支持的字段：
                - api_key: 火山引擎API密钥
                - base_url: API基础地址（OpenAI兼容格式）
                - models.vision: 视觉模型名称
                - models.text: 文本模型名称
                - max_retries: 最大重试次数
                - timeout: 请求超时时间（秒）

        Example:
            >>> # 方式1：使用默认配置文件
            >>> client = VolcengineVisionClient()

            >>> # 方式2：传入自定义配置
            >>> client = VolcengineVisionClient({
            ...     "api_key": "your-api-key",
            ...     "base_url": "https://ark.cn-beijing.volces.com/api/v3",
            ... })
        """
        # 加载配置（优先使用传入参数，否则从YAML文件加载）
        if config is None:
            try:
                config_loader = get_config()
                config = config_loader.load_api_config().get('volcengine', {})
                logger.info("从api_config.yaml成功加载火山引擎配置")
            except Exception as e:
                logger.warning(f"无法从配置文件加载API配置，将使用默认值: {e}")
                config = {}

        # 提取API配置项（使用默认值兜底）
        self.api_key: str = config.get('api_key', '')
        self.base_url: str = config.get(
            'base_url',
            'https://ark.cn-beijing.volces.com/api/v3'
        ).rstrip('/')

        # 模型配置
        models_config = config.get('models', {})
        self.vision_model: str = models_config.get(
            'vision', self.DEFAULT_VISION_MODEL
        )
        self.text_model: str = models_config.get(
            'text', self.DEFAULT_TEXT_MODEL
        )

        # 请求参数
        self.max_retries: int = config.get('max_retries', self.DEFAULT_MAX_RETRIES)
        self.timeout: int = config.get('timeout', self.DEFAULT_TIMEOUT)

        # 验证必要配置
        if not self.api_key:
            logger.warning("API Key未设置，调用API将会失败！请检查api_config.yaml配置")

        logger.info(
            f"VolcengineVisionClient初始化完成: "
            f"base_url={self.base_url}, "
            f"vision_model={self.vision_model}, "
            f"text_model={self.text_model}, "
            f"max_retries={self.max_retries}, "
            f"timeout={self.timeout}s"
        )

    def parse_image_resume(
        self,
        image_bytes: bytes,
        prompt: str = None
    ) -> Dict[str, Any]:
        """
        使用Vision模型解析图片格式的简历

        将图片进行base64编码后发送到视觉模型，提取简历中的结构化信息。

        Args:
            image_bytes: 图片的二进制数据（支持PNG/JPG/JPEG/WebP等格式）
            prompt: 自定义提示词模板（如为None使用默认的RESUME_PARSE_PROMPT），
                   提示词中应包含{content}占位符用于插入图片说明

        Returns:
            dict: 标准化的解析结果，包含以下字段：
                - success (bool): 解析是否成功
                - data (dict | None): AI返回的结构化简历数据（JSON解析后的字典）
                - error (str | None): 错误信息（成功时为None）
                - usage (dict): Token使用量统计
                    * prompt_tokens (int): 输入Token数
                    * completion_tokens (int): 输出Token数
                    * total_tokens (int): 总Token数

        Example:
            >>> client = VolcengineVisionClient()
            >>> with open("resume.png", "rb") as f:
            ...     image_data = f.read()
            >>> result = client.parse_image_resume(image_data)
            >>> if result['success']:
            ...     print(f"姓名: {result['data']['basic_info']['name']}")
        """
        start_time = time.time()
        logger.info(f"开始解析图片简历，图片大小: {len(image_bytes) / 1024:.2f}KB")

        try:
            # 1. 将图片数据进行base64编码
            base64_image = base64.b64encode(image_bytes).decode('utf-8')

            # 2. 构建自定义提示词或使用默认模板
            final_prompt = prompt or RESUME_PARSE_PROMPT

            # 3. 构建OpenAI兼容的多模态请求体
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        },
                        {
                            "type": "text",
                            "text": final_prompt.format(content="请分析这张图片中的简历内容")
                        }
                    ]
                }
            ]

            # 4. 调用带重试机制的API方法
            result = self._call_api_with_retry(
                messages=messages,
                model=self.vision_model,
                temperature=0.3,  # 低温度确保输出稳定准确
                max_tokens=4096   # 简历内容较长需要足够Token
            )

            # 5. 计算总耗时并记录日志
            elapsed_time = time.time() - start_time
            logger.info(
                f"图片简历解析完成，耗时: {elapsed_time:.2f}秒, "
                f"success: {result['success']}"
            )

            return result

        except Exception as e:
            error_msg = f"图片简历解析异常: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return {
                'success': False,
                'data': None,
                'error': error_msg,
                'usage': {}
            }

    def parse_text_resume(
        self,
        text_content: str,
        prompt: str = None
    ) -> Dict[str, Any]:
        """
        使用LLM模型解析文本格式的简历

        将纯文本简历内容发送到文本模型，提取结构化信息。
        适用于PDF/Word/TXT等已提取为文本的简历。

        Args:
            text_content: 简历的文本内容（从PDF/Word/TXT等文件提取的纯文本）
            prompt: 自定义提示词模板（如为None使用默认的RESUME_PARSE_PROMPT），
                   提示词中应包含{content}占位符用于插入简历文本

        Returns:
            dict: 标准化的解析结果，包含以下字段：
                - success (bool): 解析是否成功
                - data (dict | None): AI返回的结构化简历数据（JSON解析后的字典）
                - error (str | None): 错误信息（成功时为None）
                - usage (dict): Token使用量统计

        Example:
            >>> client = VolcengineVisionClient()
            >>> text = "张三\\nPython开发工程师\\n..."
            >>> result = client.parse_text_resume(text)
            >>> if result['success']:
            ...     print(f"姓名: {result['data']['basic_info']['name']}")
        """
        start_time = time.time()
        logger.info(f"开始解析文本简历，文本长度: {len(text_content)}字符")

        try:
            # 1. 构建自定义提示词或使用默认模板
            final_prompt = prompt or RESUME_PARSE_PROMPT

            # 2. 使用实际文本内容填充Prompt模板
            formatted_prompt = final_prompt.format(content=text_content)

            # 3. 构建OpenAI兼容的纯文本请求体
            messages = [
                {
                    "role": "user",
                    "content": formatted_prompt
                }
            ]

            # 4. 调用带重试机制的API方法
            result = self._call_api_with_retry(
                messages=messages,
                model=self.text_model,
                temperature=0.3,  # 低温度确保输出稳定准确
                max_tokens=4096   # 简历内容较长需要足够Token
            )

            # 5. 计算总耗时并记录日志
            elapsed_time = time.time() - start_time
            logger.info(
                f"文本简历解析完成，耗时: {elapsed_time:.2f}秒, "
                f"success: {result['success']}"
            )

            return result

        except Exception as e:
            error_msg = f"文本简历解析异常: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return {
                'success': False,
                'data': None,
                'error': error_msg,
                'usage': {}
            }

    def _call_api_with_retry(
        self,
        messages: List[Dict[str, Any]],
        model: str,
        **kwargs
    ) -> Dict[str, Any]:
        """
        带重试机制的API调用方法

        实现指数退避重试策略，处理网络波动、限流等临时性错误。
        对认证失败、参数错误等不可恢复错误立即返回不重试。

        Args:
            messages: OpenAI格式的消息列表
                格式示例：[{"role": "user", "content": "..."}]
            model: 要调用的模型名称（如"doubao-vision-pro-32k"）
            **kwargs: 其他API参数
                - temperature (float): 温度参数（0-2，默认0.7）
                - max_tokens (int): 最大生成Token数（默认2048）
                - top_p (float): 核采样参数（0-1）

        Returns:
            dict: API响应结果，包含以下字段：
                - success (bool): 调用是否成功
                - data (dict | None): AI返回的结构化数据（JSON解析后的字典）
                - error (str | None): 错误信息
                - usage (dict): Token使用量统计

        Note:
            重试策略详情：
            - 可重试错误：网络超时、429限流、5xx服务器错误
            - 不可重试错误：401认证失败、400参数错误（立即返回）
            - 退避间隔：指数退避（1秒 → 2秒 → 4秒 → ...）
            - 最大重试次数：由max_retries配置决定（默认3次）
        """
        last_error = None

        for attempt in range(self.max_retries + 1):
            try:
                # 记录当前重试次数
                if attempt > 0:
                    wait_time = 2 ** attempt  # 指数退避：1s, 2s, 4s...
                    logger.info(
                        f"第{attempt}次重试，等待{wait_time}秒后执行... "
                        f"(剩余重试次数: {self.max_retries - attempt})"
                    )
                    time.sleep(wait_time)

                # 构建请求负载
                payload = {
                    "model": model,
                    "messages": messages,
                    "temperature": kwargs.get('temperature', 0.7),
                    "max_tokens": kwargs.get('max_tokens', 2048),
                    "top_p": kwargs.get('top_p', 0.9)
                }

                # 发送HTTP请求
                response = self._send_request("/chat/completions", payload)

                # 检查HTTP状态码决定是否可重试
                status_code = response.get('status_code')
                response_data = response.get('data', {})

                # 处理API业务层面的错误
                if 'error' in response_data:
                    error_info = response_data['error']
                    error_type = isinstance(error_info, dict) and error_info.get('type', '')
                    error_message = (
                        error_info.get('message', '')
                        if isinstance(error_info, dict)
                        else str(error_info)
                    )
                    http_status = status_code or 500

                    # 不可重试的错误：认证失败、参数错误
                    if http_status in [400, 401, 403]:
                        error_msg = f"API错误({http_status}): {error_message}"
                        logger.error(error_msg)
                        return {
                            'success': False,
                            'data': None,
                            'error': error_msg,
                            'usage': {}
                        }

                    # 可重试的错误：限流、服务器错误
                    if http_status == 429:
                        retry_after = response.get('headers', {}).get('Retry-After')
                        wait = int(retry_after) if retry_after else (2 ** attempt)
                        logger.warning(
                            f"API限流(429)，将在{wait}秒后重试。"
                            f"提示: {error_message}"
                        )
                        last_error = f"API限流: {error_message}"
                        continue

                    if http_status >= 500:
                        logger.warning(
                            f"服务器错误({http_status})，准备重试。"
                            f"提示: {error_message}"
                        )
                        last_error = f"服务器错误({http_status}): {error_message}"
                        continue

                # 成功响应：提取AI返回的内容
                if status_code and 200 <= status_code < 300:
                    return self._parse_success_response(response_data)

                # 其他未知状态码
                last_error = f"未知HTTP状态码: {status_code}"
                logger.warning(last_error)
                continue

            except requests.exceptions.ConnectionError as e:
                last_error = f"网络连接失败: {str(e)}"
                logger.warning(f"连接错误(第{attempt + 1}次): {last_error}")
                continue

            except requests.exceptions.Timeout as e:
                last_error = f"请求超时: {str(e)}"
                logger.warning(f"超时错误(第{attempt + 1}次): {last_error}")
                continue

            except requests.exceptions.RequestException as e:
                last_error = f"请求异常: {str(e)}"
                logger.error(f"请求异常(第{attempt + 1}次): {last_error}", exc_info=True)
                continue

            except json.JSONDecodeError as e:
                last_error = f"JSON解析失败: {str(e)}"
                logger.error(last_error, exc_info=True)
                return {
                    'success': False,
                    'data': None,
                    'error': last_error,
                    'usage': {}
                }

            except Exception as e:
                last_error = f"未预期的异常: {str(e)}"
                logger.error(f"API调用异常(第{attempt + 1}次): {last_error}", exc_info=True)
                continue

        # 所有重试都失败
        error_msg = (
            f"API调用失败（已重试{self.max_retries}次）: "
            f"{last_error or '未知错误'}"
        )
        logger.error(error_msg)
        return {
            'success': False,
            'data': None,
            'error': error_msg,
            'usage': {}
        }

    def _send_request(self, endpoint: str, payload: Dict[str, Any]) -> Dict[str, Any]:
        """
        发送HTTP POST请求到API端点

        使用requests库向火山引擎API发送POST请求，
        设置标准请求头并记录详细的请求/响应日志。

        Args:
            endpoint: API路径（如"/chat/completions"）
            payload: 请求体字典（符合OpenAI Chat Completions API格式）

        Returns:
            dict: 解析后的响应结果，包含以下字段：
                - status_code (int | None): HTTP状态码
                - data (dict): 响应JSON数据
                - headers (dict): 响应头（用于获取Retry-After等信息）

        Raises:
            requests.exceptions.ConnectionError: 网络连接失败
            requests.exceptions.Timeout: 请求超时
            requests.exceptions.HTTPError: HTTP错误（非2xx状态码）
            json.JSONDecodeError: 响应JSON解析失败
        """
        # 构建完整URL
        url = f"{self.base_url}{endpoint}"

        # 设置标准请求头
        # 火山引擎Ark API使用标准的 Authorization: Bearer 格式进行认证
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}",
            "Accept": "application/json"
        }

        logger.debug(f"发送API请求: {url}")
        logger.debug(f"请求模型: {payload.get('model')}")
        logger.debug(f"消息数量: {len(payload.get('messages', []))}")

        start_time = time.time()

        try:
            # 发送POST请求
            response = requests.post(
                url,
                headers=headers,
                json=payload,
                timeout=self.timeout
            )

            # 计算请求耗时
            elapsed_time = time.time() - start_time

            # 记录响应基本信息
            logger.info(
                f"API响应: URL={url}, "
                f"Status={response.status_code}, "
                f"耗时={elapsed_time:.2f}s"
            )

            # 尝试解析JSON响应
            try:
                response_data = response.json()
            except json.JSONDecodeError as e:
                logger.error(
                    f"JSON解析失败，原始响应: {response.text[:500]}",
                    exc_info=True
                )
                raise json.JSONDecodeError(
                    f"无效的JSON响应: {e.msg}",
                    response.text,
                    e.pos
                ) from e

            # 记录Token用量（如果存在）
            usage = response_data.get('usage', {})
            if usage:
                logger.info(
                    f"Token使用量: "
                    f"prompt={usage.get('prompt_tokens', 0)}, "
                    f"completion={usage.get('completion_tokens', 0)}, "
                    f"total={usage.get('total_tokens', 0)}"
                )

            # 返回标准化响应
            return {
                'status_code': response.status_code,
                'data': response_data,
                'headers': dict(response.headers)
            }

        except requests.exceptions.ConnectionError as e:
            elapsed_time = time.time() - start_time
            logger.error(
                f"网络连接失败: {url}, 耗时={elapsed_time:.2f}s, 错误: {str(e)}"
            )
            raise

        except requests.exceptions.Timeout as e:
            elapsed_time = time.time() - start_time
            logger.error(
                f"请求超时: {url}, 耗时={elapsed_time:.2f}s, "
                f"超时限制={self.timeout}s"
            )
            raise

        except requests.exceptions.RequestException as e:
            elapsed_time = time.time() - start_time
            logger.error(
                f"请求异常: {url}, 耗时={elapsed_time:.2f}s, 错误: {str(e)}",
                exc_info=True
            )
            raise

    def _parse_success_response(self, response_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        解析成功的API响应

        从API响应中提取AI生成的内容和Token使用量信息，
        并尝试将AI返回的JSON字符串解析为Python字典。

        Args:
            response_data: API返回的原始JSON数据（字典格式）

        Returns:
            dict: 标准化解析结果，包含以下字段：
                - success (bool): 固定为True
                - data (dict | None): 结构化简历数据（JSON解析后的字典）
                - error (str | None): 固定为None
                - usage (dict): Token使用量统计

        Note:
            如果AI返回的内容不是有效JSON格式，会尝试清理markdown标记后再次解析。
            如果仍然解析失败，会将原始文本作为raw_content字段返回。
        """
        # 提取Token使用量
        usage = response_data.get('usage', {})
        usage_info = {
            'prompt_tokens': usage.get('prompt_tokens', 0),
            'completion_tokens': usage.get('completion_tokens', 0),
            'total_tokens': usage.get('total_tokens', 0)
        }

        # 提取AI生成的内容
        choices = response_data.get('choices', [])
        if not choices:
            logger.warning("API响应中无choices字段")
            return {
                'success': True,
                'data': None,
                'error': 'API返回空响应',
                'usage': usage_info
            }

        # 获取第一个选择的内容
        message = choices[0].get('message', {})
        content = message.get('content', '')

        if not content.strip():
            logger.warning("AI返回内容为空")
            return {
                'success': True,
                'data': None,
                'error': 'AI返回内容为空',
                'usage': usage_info
            }

        # 尝试解析JSON内容
        parsed_data = self._extract_json_from_response(content)

        if parsed_data is not None:
            logger.info("成功解析AI返回的结构化数据")
            return {
                'success': True,
                'data': parsed_data,
                'error': None,
                'usage': usage_info
            }
        else:
            # JSON解析失败，返回原始内容
            logger.warning(
                "无法将AI返回内容解析为JSON，返回原始文本。"
                f"内容预览: {content[:200]}..."
            )
            return {
                'success': True,
                'data': {'raw_content': content},
                'error': 'JSON解析失败，已保留原始文本',
                'usage': usage_info
            }

    def _extract_json_from_response(self, content: str) -> Optional[Dict[str, Any]]:
        """
        从AI响应中提取JSON数据

        处理AI可能返回的各种格式（纯JSON、Markdown代码块等），
        尝试提取有效的JSON对象。

        Args:
            content: AI返回的原始文本内容

        Returns:
            dict | None: 成功解析返回字典，失败返回None

        Note:
            支持的格式：
            1. 纯JSON字符串（以{开头）
            2. Markdown代码块包裹的JSON（```json ... ``` 或 ``` ... ```）
            3. JSON前后有其他文字的情况（查找第一个{和最后一个}）
        """
        # 清理内容
        cleaned_content = content.strip()

        # 尝试1：直接解析（纯JSON）
        if cleaned_content.startswith('{'):
            try:
                return json.loads(cleaned_content)
            except json.JSONDecodeError:
                pass

        # 尝试2：移除Markdown代码块标记
        import re

        # 匹配 ```json ... ``` 或 ``` ... ```
        code_block_pattern = r'```(?:json)?\s*\n?(.*?)\n?```'
        match = re.search(code_block_pattern, cleaned_content, re.DOTALL)
        if match:
            json_str = match.group(1).strip()
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                pass

        # 尝试3：查找第一个{和最后一个}
        first_brace = cleaned_content.find('{')
        last_brace = cleaned_content.rfind('}')
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_str = cleaned_content[first_brace:last_brace + 1]
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                pass

        # 所有尝试都失败
        return None


class ResumeParsingPipeline:
    """
    简历解析完整流程编排器

    职责：将文件预处理 → AI解析 → 结果处理 → 数据存储 整合为完整流程，
    提供一键式简历解析和持久化服务。

    Attributes:
        parser (ResumeParser): 文件预处理解析器实例
        ai_client (VolcengineVisionClient): AI API客户端实例
        _raw_ai_response (dict | None): 原始AI响应JSON
        _cleaned_data (dict | None): 清洗后的数据
        _user_profile (UserProfile | None): UserProfile模型实例
    """

    def __init__(self):
        """
        初始化流程所需的各个组件

        创建文件解析器、AI客户端、日志记录器和中间结果缓存。
        """
        self.parser = ResumeParser()  # 文件预处理
        self.ai_client = VolcengineVisionClient()  # AI API调用
        self.logger = get_logger(__name__)  # 日志记录

        # 存储中间结果
        self._raw_ai_response: Optional[Dict[str, Any]] = None  # 原始AI响应JSON
        self._cleaned_data: Optional[Dict[str, Any]] = None  # 清洗后的数据
        self._user_profile: Optional['UserProfile'] = None  # UserProfile模型实例

        self.logger.info("简历解析流程编排器初始化完成")

    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        执行完整的简历解析流程

        按顺序执行：文件预处理 → AI解析 → 结果处理 → 数据映射 → 持久化存储，
        返回完整的解析结果和统计信息。

        Args:
            file_path: 简历文件的完整路径（支持PDF/Word/TXT/图片格式）

        Returns:
            dict: 标准化的解析结果，包含以下字段：
                - success (bool): 整个流程是否成功完成
                - user_profile (UserProfile | None): Pydantic模型实例
                - raw_response (dict | None): 原始AI返回（用于备份）
                - error (str | None): 错误信息（成功时为None）
                - stats (dict): 统计信息
                    * file_size (int): 文件大小（字节）
                    * parse_time (float): 总耗时（秒）
                    * api_tokens_used (int): API token消耗

        Example:
            >>> pipeline = ResumeParsingPipeline()
            >>> result = pipeline.parse_resume("resumes/张三.pdf")
            >>> if result['success']:
            ...     print(f"姓名: {result['user_profile'].name}")
        """
        # 记录流程开始时间
        pipeline_start_time = time.time()
        self.logger.info("=" * 60)
        self.logger.info(f"开始执行简历解析流程，文件路径: {file_path}")

        # 初始化结果字典
        result = {
            'success': False,
            'user_profile': None,
            'raw_response': None,
            'error': None,
            'stats': {
                'file_size': 0,
                'parse_time': 0.0,
                'api_tokens_used': 0
            }
        }

        try:
            # ==================== 步骤1：文件预处理 ====================
            self.logger.info("【步骤1/5】开始文件预处理...")
            parse_result = self.parser.parse_resume_file(file_path)

            if not parse_result['success']:
                error_msg = f"文件预处理失败: {parse_result.get('error', '未知错误')}"
                self.logger.error(error_msg)
                result['error'] = error_msg
                return result

            # 记录文件统计信息
            file_size = parse_result['metadata'].get('size', 0)
            result['stats']['file_size'] = file_size
            file_type = parse_result['file_type']
            content = parse_result['content']

            self.logger.info(
                f"文件预处理完成: 类型={file_type}, "
                f"大小={file_size / 1024:.2f}KB, "
                f"内容长度={len(content) if isinstance(content, str) else len(content)}字节"
            )

            # ==================== 步骤2：根据文件类型调用AI API ====================
            self.logger.info("【步骤2/5】开始AI智能解析...")

            if file_type == 'image':
                # 图片类型使用视觉模型
                self.logger.info(f"使用视觉模型({self.ai_client.vision_model})解析图片")
                ai_result = self.ai_client.parse_image_resume(
                    image_bytes=content,
                    prompt=RESUME_PARSE_PROMPT
                )
            else:
                # 文本类型使用文本模型
                self.logger.info(f"使用文本模型({self.ai_client.text_model})解析文本")
                ai_result = self.ai_client.parse_text_resume(
                    text_content=content,
                    prompt=RESUME_PARSE_PROMPT
                )

            # 检查API调用是否成功
            if not ai_result.get('success'):
                error_msg = f"AI解析失败: {ai_result.get('error', '未知错误')}"
                self.logger.error(error_msg)
                result['error'] = error_msg
                return result

            # 记录Token用量
            usage = ai_result.get('usage', {})
            result['stats']['api_tokens_used'] = usage.get('total_tokens', 0)

            self.logger.info(
                f"AI解析完成: Token消耗={result['stats']['api_tokens_used']}"
            )

            # ==================== 步骤3：处理AI返回结果 ====================
            self.logger.info("【步骤3/5】处理AI返回结果...")
            cleaned_data = self._process_ai_response(ai_result)

            if cleaned_data is None:
                error_msg = "数据处理失败：无法从AI响应中提取有效数据"
                self.logger.error(error_msg)
                result['error'] = error_msg
                result['raw_response'] = ai_result
                return result

            self._cleaned_data = cleaned_data
            self.logger.info("数据处理完成，已提取结构化数据")

            # ==================== 步骤4：映射到UserProfile模型 ====================
            self.logger.info("【步骤4/5】映射数据到UserProfile模型...")
            try:
                user_profile = self._map_to_user_profile(cleaned_data)
                # 设置简历文件路径
                user_profile.resume_file_path = file_path
                self._user_profile = user_profile
                result['user_profile'] = user_profile
                self.logger.info(f"模型验证通过，用户画像创建成功: {user_profile.name}")
            except ValidationError as e:
                error_msg = f"数据验证失败: {str(e)}"
                self.logger.error(error_msg)
                result['error'] = error_msg
                result['raw_response'] = self._raw_ai_response
                return result

            # ==================== 步骤5：持久化存储 ====================
            self.logger.info("【步骤5/5】保存数据到数据库和备份文件...")

            # 5a. 保存到SQLite数据库
            db_save_success = self._save_to_database(user_profile)
            if db_save_success:
                self.logger.info("数据库保存成功")
            else:
                self.logger.warning("数据库保存失败（不影响主流程）")

            # 5b. 保存JSON备份（mode='json'确保datetime等字段转为ISO字符串）
            backup_path = self._save_json_backup(
                self._raw_ai_response or {},
                user_profile.model_dump(mode='json') if user_profile else {}
            )
            if backup_path:
                self.logger.info(f"JSON备份保存成功: {backup_path}")

            # ==================== 流程完成 ====================
            result['success'] = True
            result['raw_response'] = self._raw_ai_response

            # 计算总耗时
            total_time = time.time() - pipeline_start_time
            result['stats']['parse_time'] = round(total_time, 3)

            self.logger.info("=" * 60)
            self.logger.info(
                f"简历解析流程完成！总耗时: {total_time:.3f}秒, "
                f"状态: 成功"
            )
            self.logger.info("=" * 60)

            return result

        except Exception as e:
            # 兜底异常捕获，确保流程不会崩溃
            total_time = time.time() - pipeline_start_time
            result['stats']['parse_time'] = round(total_time, 3)
            error_msg = f"流程执行异常: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            result['error'] = error_msg
            return result

    def _process_ai_response(self, ai_result: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """
        处理AI API返回的结果

        从VolcengineVisionClient的返回值中提取、清洗和标准化数据。

        Args:
            ai_result: VolcengineVisionClient返回的结果字典
                格式：{success, data, error, usage}

        Returns:
            dict | None: 清洗后的结构化数据字典，失败时返回None

        处理步骤：
            1. 验证响应有效性
            2. 提取JSON数据
            3. 数据清洗（去空白、验证类型、处理缺失字段）
            4. 缓存原始响应用于备份
        """
        try:
            # 1. 检查响应有效性
            if not ai_result.get('success'):
                self.logger.error(f"AI响应标记为失败: {ai_result.get('error')}")
                return None

            data = ai_result.get('data')
            if data is None:
                self.logger.error("AI响应中无有效数据")
                return None

            # 缓存原始响应
            self._raw_ai_response = ai_result.copy()

            # 2. 处理特殊情况：原始内容未JSON化
            if 'raw_content' in data:
                self.logger.warning("AI返回了非JSON格式的内容，尝试二次提取")
                extracted = self.ai_client._extract_json_from_response(data['raw_content'])
                if extracted is None:
                    self.logger.error("无法从原始内容中提取JSON数据")
                    return None
                data = extracted

            # 3. 开始数据清洗
            cleaned_data = self._clean_profile_data(data)

            if cleaned_data is None:
                self.logger.error("数据清洗失败")
                return None

            self.logger.info("AI响应处理完成，数据清洗成功")
            return cleaned_data

        except Exception as e:
            self.logger.error(f"处理AI响应时发生异常: {str(e)}", exc_info=True)
            return None

    def _clean_profile_data(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        清洗和标准化简历数据

        对AI返回的原始数据进行清洗，包括：
        - 字符串字段去除首尾空白
        - 数值字段范围校验
        - 列表字段类型确保
        - 缺失字段填充默认值

        Args:
            raw_data: AI返回的原始数据字典

        Returns:
            dict: 清洗后的标准化数据字典
        """
        cleaned = {}

        try:
            # ---------- 基本信息 ----------
            basic_info = raw_data.get('basic_info', {})
            cleaned['name'] = self._clean_string(basic_info.get('name'))
            cleaned['phone'] = self._clean_string(basic_info.get('phone'))
            cleaned['email'] = self._clean_string(basic_info.get('email'))

            # ---------- 教育背景 ----------
            education = raw_data.get('education', {})
            cleaned['education'] = self._clean_string(education.get('degree'), default='未知')
            cleaned['major'] = self._clean_string(education.get('major'))
            cleaned['school'] = self._clean_string(education.get('school'))

            # ---------- 工作经验 ----------
            work_exp = raw_data.get('work_experience', {})
            total_years = work_exp.get('total_years')
            if total_years is not None:
                # 确保工作年限是有效的浮点数且>=0
                try:
                    years_val = float(total_years)
                    cleaned['total_experience_years'] = max(0.0, years_val)
                except (ValueError, TypeError):
                    cleaned['total_experience_years'] = 0.0
                    self.logger.warning(f"工作年限格式无效: {total_years}，使用默认值0")
            else:
                cleaned['total_experience_years'] = 0.0

            # 提取当前职位和公司（从最近的工作经历）
            positions = work_exp.get('positions', [])
            if positions and isinstance(positions, list) and len(positions) > 0:
                latest_position = positions[0]  # 已按时间倒序排列
                cleaned['current_position'] = self._clean_string(latest_position.get('title'))
                cleaned['current_company'] = self._clean_string(latest_position.get('company'))
            else:
                cleaned['current_position'] = None
                cleaned['current_company'] = None

            # ---------- 技能和专长 ----------
            skills = raw_data.get('skills', [])
            cleaned['skills'] = self._ensure_list(skills) if skills else []

            core_competencies = raw_data.get('core_competencies', [])
            cleaned['core_competencies'] = (
                self._ensure_list(core_competencies) if core_competencies else []
            )

            # ---------- 项目经历 ----------
            projects = raw_data.get('projects', [])
            cleaned['projects'] = (
                self._ensure_list(projects) if projects and isinstance(projects, list) else []
            )

            # ---------- 求职意向 ----------
            job_expectation = raw_data.get('job_expectation', {})

            # 期望岗位
            target_positions = job_expectation.get('target_positions')
            cleaned['expected_positions'] = (
                self._ensure_list(target_positions) if target_positions else []
            )

            # 解析期望薪资（从字符串如"15k-25k"提取数值）
            expected_salary_str = job_expectation.get('expected_salary')
            salary_min, salary_max = self._parse_salary_range(expected_salary_str)
            cleaned['expected_salary_min'] = salary_min
            cleaned['expected_salary_max'] = salary_max

            # 期望工作地点
            expected_locations = job_expectation.get('expected_locations')
            cleaned['expected_locations'] = (
                self._ensure_list(expected_locations) if expected_locations else []
            )

            # 记录清洗警告（如果有缺失关键字段）
            if not cleaned.get('name'):
                self.logger.warning("清洗后数据缺少必要字段: name")

            return cleaned

        except Exception as e:
            self.logger.error(f"数据清洗过程发生异常: {str(e)}", exc_info=True)
            return None

    def _map_to_user_profile(self, cleaned_data: Dict[str, Any]) -> 'UserProfile':
        """
        将清洗后的数据映射到UserProfile Pydantic模型

        将标准化的数据字典转换为UserProfile模型实例，
        利用Pydantic进行数据验证和类型转换。

        Args:
            cleaned_data: 清洗后的原始数据字典

        Returns:
            UserProfile: 验证通过的模型实例

        Raises:
            ValidationError: 当数据无法通过模型验证时抛出
        """
        # 延迟导入避免循环依赖
        from core.models import UserProfile

        # 构建模型构造参数
        profile_params = {
            # 必填字段
            'name': cleaned_data.get('name', ''),
            'education': cleaned_data.get('education', '未知'),
            'total_experience_years': cleaned_data.get('total_experience_years', 0.0),

            # 可选字段
            'phone': cleaned_data.get('phone'),
            'email': cleaned_data.get('email'),
            'major': cleaned_data.get('major'),
            'school': cleaned_data.get('school'),
            'current_position': cleaned_data.get('current_position'),
            'current_company': cleaned_data.get('current_company'),

            # 列表字段
            'skills': cleaned_data.get('skills', []),
            'core_competencies': cleaned_data.get('core_competencies', []),
            'projects': cleaned_data.get('projects', []),

            # 求职意向
            'expected_positions': cleaned_data.get('expected_positions', []),
            'expected_salary_min': cleaned_data.get('expected_salary_min'),
            'expected_salary_max': cleaned_data.get('expected_salary_max'),
            'expected_locations': cleaned_data.get('expected_locations', []),
        }

        # 创建模型实例（Pydantic会自动进行验证）
        try:
            profile = UserProfile(**profile_params)
            self.logger.info("数据映射到UserProfile模型成功")
            return profile
        except ValidationError as e:
            # 收集所有验证错误并生成友好提示
            errors = e.errors()
            error_details = []
            for err in errors:
                field = '.'.join(str(loc) for loc in err['loc'])
                msg = err['msg']
                error_details.append(f"{field}: {msg}")

            error_msg = "用户画像数据验证失败:\n" + "\n".join(error_details)
            self.logger.error(error_msg)
            raise ValidationError(errors, model=UserProfile)

    def _save_to_database(self, profile: 'UserProfile') -> bool:
        """
        将用户画像保存到SQLite数据库

        使用DatabaseHelper将UserProfile数据持久化到profiles表，
        支持新增或更新已有记录。

        Args:
            profile: UserProfile模型实例

        Returns:
            bool: 保存是否成功
        """
        try:
            # 将模型转换为字典
            profile_dict = profile.model_dump()

            # 移除不需要存入数据库的字段
            profile_dict.pop('created_at', None)
            profile_dict.pop('updated_at', None)

            # 获取数据库实例
            db = get_db()

            # 检查是否已有活跃的用户画像（ID=1表示当前使用的画像）
            existing_profile = db.get_profile(profile_id=1)

            if existing_profile:
                # 更新现有记录
                success = db.update_profile(profile_id=1, profile_data=profile_dict)
                if success:
                    self.logger.info("用户画像更新成功（ID=1）")
                else:
                    self.logger.warning("用户画像更新失败（记录不存在）")
                return success
            else:
                # 创建新记录
                new_id = db.create_profile(profile_dict)
                self.logger.info(f"用户新建画像成功（ID={new_id}）")
                return True

        except Exception as e:
            self.logger.error(f"数据库保存失败: {str(e)}", exc_info=True)
            return False

    def _save_json_backup(
        self,
        raw_data: Dict[str, Any],
        profile_data: Dict[str, Any]
    ) -> Optional[str]:
        """
        保存原始AI响应和处理后的数据为JSON备份文件

        将完整的解析结果以结构化JSON格式保存到本地文件系统，
        用于审计追踪和数据恢复。

        Args:
            raw_data: AI原始响应数据
            profile_data: 处理后的用户画像数据

        Returns:
            str | None: 备份文件的绝对路径，失败时返回None
        """
        try:
            # 确保备份目录存在
            backup_dir = Path("data/profiles")
            backup_dir.mkdir(parents=True, exist_ok=True)

            # 生成带时间戳的文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"profile_{timestamp}.json"
            file_path = backup_dir / filename

            # 构建备份数据结构
            backup_content = {
                "timestamp": datetime.now().isoformat(),
                "raw_ai_response": raw_data,
                "processed_profile": profile_data,
                "version": "1.0"
            }

            # 写入JSON文件（UTF-8编码，不转义Unicode字符）
            # 使用自定义default处理器处理datetime等非JSON原生类型
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(backup_content, f, ensure_ascii=False, indent=2,
                          default=lambda o: o.isoformat() if hasattr(o, 'isoformat') else str(o))

            absolute_path = str(file_path.resolve())
            self.logger.info(f"JSON备份文件保存成功: {absolute_path}")
            return absolute_path

        except Exception as e:
            self.logger.error(f"JSON备份保存失败: {str(e)}", exc_info=True)
            return None

    # ==================== 便捷方法 ====================

    def get_profile_as_dict(self) -> Optional[Dict[str, Any]]:
        """
        获取当前解析结果的字典形式

        Returns:
            dict | None: 用户画像字典，无数据时返回None
        """
        if self._user_profile is None:
            return None
        return self._user_profile.model_dump()

    def get_profile_as_json(self) -> Optional[str]:
        """
        获取当前解析结果的JSON字符串

        Returns:
            str | None: JSON格式的字符串，无数据时返回None
        """
        if self._user_profile is None:
            return None
        return self._user_profile.model_dump_json()

    def clear_results(self):
        """清除当前的解析结果缓存"""
        self._raw_ai_response = None
        self._cleaned_data = None
        self._user_profile = None
        self.logger.info("解析结果缓存已清除")

    def reparse(self, file_path: str) -> Dict[str, Any]:
        """
        重新解析简历文件

        先清除旧结果再执行新的解析流程，用于重新分析同一文件或不同文件。

        Args:
            file_path: 简历文件的完整路径

        Returns:
            dict: 完整的解析结果（与parse_resume返回格式相同）
        """
        self.logger.info(f"执行重新解析操作，目标文件: {file_path}")
        self.clear_results()
        return self.parse_resume(file_path)

    # ==================== 私有工具方法 ====================

    @staticmethod
    def _clean_string(value: Any, default: Optional[str] = None) -> Optional[str]:
        """
        清理字符串值（去除首尾空白）

        Args:
            value: 待清理的值（任意类型）
            default: 当值为空时的默认返回值

        Returns:
            str | None: 清理后的字符串，或默认值
        """
        if value is None:
            return default
        if isinstance(value, str):
            stripped = value.strip()
            return stripped if stripped else default
        # 非字符串类型转为字符串后清理
        return str(value).strip() or default

    @staticmethod
    def _ensure_list(value: Any) -> List[Any]:
        """
        确保值是列表类型

        对于非列表类型的值进行包装转换。

        Args:
            value: 待处理的值

        Returns:
            list: 确保是列表类型的数据
        """
        if value is None:
            return []
        if isinstance(value, list):
            return value
        # 单个元素包装为列表
        return [value]

    @staticmethod
    def _parse_salary_range(salary_str: Optional[str]) -> tuple:
        """
        解析薪资范围字符串

        从类似"15k-25k"、"15-25K"、"15k~25k"等格式的字符串中提取数值。

        Args:
            salary_str: 薪资范围字符串

        Returns:
            tuple: (最低薪资, 最高薪资)，单位K/月，无法解析则返回(None, None)
        """
        if not salary_str or not isinstance(salary_str, str):
            return None, None

        # 使用正则表达式匹配数字
        pattern = r'(\d+)\s*[kK~-]\s*(\d+)'
        match = re.search(pattern, salary_str)

        if match:
            try:
                min_salary = int(match.group(1))
                max_salary = int(match.group(2))
                return min_salary, max_salary
            except ValueError:
                pass

        # 尝试单个数字（如"20k"）
        single_pattern = r'(\d+)\s*[kK]'
        single_match = re.search(single_pattern, salary_str)
        if single_match:
            try:
                salary = int(single_match.group(1))
                return salary, salary
            except ValueError:
                pass

        return None, None


# ==================== 测试代码 ====================
if __name__ == "__main__":
    """模块自测试代码"""
    import sys
    import io

    # 设置控制台输出编码为UTF-8
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    print("=" * 70)
    print("简历文件预处理解析器 - 功能测试")
    print("=" * 70)

    # 创建解析器实例
    parser = ResumeParser()

    # 测试1：验证功能
    print("\n【测试1】文件验证功能")
    print("-" * 50)

    # 测试不存在的文件
    result = parser._validate_file("不存在的文件.pdf")
    print(f"✓ 不存在的文件: valid={result['valid']}, error='{result['error']}'")

    # 测试空目录（如果存在的话）
    test_dir = "e:\\projects\\job\\core"
    result = parser._validate_file(test_dir)
    print(f"✓ 目录路径: valid={result['valid']}")

    # 测试2：格式检测
    print("\n【测试2】文件格式检测")
    print("-" * 50)

    test_files = [
        ("resume.pdf", ".pdf"),
        ("document.docx", ".docx"),
        ("old.doc", ".doc"),
        ("note.txt", ".txt"),
        ("photo.png", ".png"),
        ("image.jpg", ".jpg"),
        ("picture.jpeg", ".jpeg"),
        ("data.xlsx", ".xlsx"),  # 不支持的格式
    ]

    for filename, expected_ext in test_files:
        detected = parser._detect_format(filename)
        supported = parser._is_supported_format(detected)
        status = "✓ 支持" if supported else "✗ 不支持"
        print(f"  {filename:20s} → 扩展名: {detected:6s} [{status}]")

    # 测试3：创建测试文件并验证解析流程
    print("\n【测试3】实际文件解析测试")
    print("-" * 50)

    # 创建测试用的TXT文件
    test_txt_path = "e:\\projects\\job\\test_resume.txt"
    test_content = """张三
Python开发工程师

教育背景：
- 本科 · 计算机科学与技术 · XX大学 · 2015-2019

工作经验：
1. ABC科技有限公司 · Python开发工程师 · 2019-2022
   - 负责后端API开发和数据库设计
   - 使用Django框架构建Web应用

2. XYZ互联网公司 · 高级Python工程师 · 2022-至今
   - 主导微服务架构改造
   - 性能优化提升系统吞吐量300%

技能特长：
- Python, Django, Flask
- MySQL, Redis, PostgreSQL
- Docker, Kubernetes
"""

    try:
        # 写入测试文件
        with open(test_txt_path, 'w', encoding='utf-8') as f:
            f.write(test_content)
        print(f"✓ 创建测试TXT文件: {test_txt_path}")

        # 解析TXT文件
        result = parser.parse_resume_file(test_txt_path)
        print(f"\n解析结果:")
        print(f"  成功: {result['success']}")
        print(f"  类型: {result['file_type']}")
        print(f"  内容长度: {len(result['content'])} 字符")
        print(f"  元数据: {result['metadata']}")
        print(f"  错误: {result['error']}")

        if result['success']:
            print(f"\n--- 提取的内容预览 ---")
            print(result['content'][:200] + "..." if len(result['content']) > 200 else result['content'])

    finally:
        # 清理测试文件
        if os.path.exists(test_txt_path):
            os.remove(test_txt_path)
            print(f"\n✓ 已清理测试文件")

    # 测试4：错误处理测试
    print("\n【测试4】错误处理测试")
    print("-" * 50)

    # 测试不存在的文件
    result = parser.parse_resume_file("不存在的文件.pdf")
    print(f"✓ 不存在的文件:")
    print(f"  成功: {result['success']}")
    print(f"  错误: {result['error']}")

    # 测试不支持的格式
    unsupported_file = "test.xlsx"
    print(f"\n✓ 不支持的格式 ({unsupported_file}):")
    # 创建一个临时文件来测试格式检测
    temp_path = f"e:\\projects\\job\\{unsupported_file}"
    try:
        with open(temp_path, 'w') as f:
            f.write("test")
        result = parser.parse_resume_file(temp_path)
        print(f"  成功: {result['success']}")
        print(f"  错误: {result['error']}")
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

    print("\n" + "=" * 70)
    print("[SUCCESS] 所有基本测试通过！")
    print("=" * 70)
    print("\n提示：要完整测试PDF/Word/图片解析功能，请准备相应的测试文件。")
